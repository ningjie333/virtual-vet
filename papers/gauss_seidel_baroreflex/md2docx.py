#!/usr/bin/env python3
"""Convert manuscript_v6.md to a Word document formatted for SIMULATION journal."""
import re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

PAPER_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(PAPER_DIR, "manuscript_v6.md")
DOCX_PATH = os.path.join(PAPER_DIR, "manuscript_v6.docx")

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Styles ──────────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0

for level in range(1, 4):
    hstyle = doc.styles[f'Heading {level}']
    hstyle.font.name = 'Times New Roman'
    hstyle.font.size = Pt(14 if level == 1 else 12)
    hstyle.font.bold = True
    hstyle.font.color.rgb = RGBColor(0, 0, 0)
    hstyle.paragraph_format.space_before = Pt(12)
    hstyle.paragraph_format.space_after = Pt(6)

# ── Read markdown ───────────────────────────────────────────────────────────────
with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# ── Parse and render ────────────────────────────────────────────────────────────
in_code_block = False
code_lines = []
in_table = False
table_rows = []
prev_was_hr = False  # track consecutive --- for page breaks

def flush_code():
    global code_lines
    if code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        code_lines = []

def flush_table():
    global table_rows, in_table
    if table_rows:
        # Parse markdown table
        data_rows = []
        for row in table_rows:
            cells = [c.strip() for c in row.strip('|').split('|')]
            data_rows.append(cells)
        if len(data_rows) >= 2:
            # Skip separator row (---|---|---)
            header = data_rows[0]
            body = [r for r in data_rows[2:] if not all(set(c.strip()) <= {'-', ':', ' '} for c in r)]
            ncols = len(header)
            table = doc.add_table(rows=1 + len(body), cols=ncols)
            table.style = 'Table Grid'
            # Header
            for i, h in enumerate(header):
                if i < ncols:
                    cell = table.rows[0].cells[i]
                    cell.text = h.strip('*').strip()
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(10)
            # Body
            for ri, row_data in enumerate(body):
                for ci, val in enumerate(row_data):
                    if ci < ncols:
                        table.rows[ri + 1].cells[ci].text = val.strip('*').strip()
                        for p in table.rows[ri + 1].cells[ci].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(10)
        table_rows = []
    in_table = False

for line in lines:
    stripped = line.rstrip('\n')

    # Code block toggle
    if stripped.startswith('```'):
        if in_code_block:
            flush_code()
            in_code_block = False
        else:
            in_code_block = True
        continue

    if in_code_block:
        code_lines.append(stripped)
        continue

    # Table detection
    if '|' in stripped and stripped.strip().startswith('|'):
        if not in_table:
            in_table = True
        table_rows.append(stripped)
        continue
    elif in_table:
        flush_table()

    # Empty line
    if not stripped.strip():
        continue

    # Horizontal rule — consecutive --- with only whitespace between = page break
    if stripped.strip() == '---':
        if prev_was_hr:
            doc.add_page_break()
            prev_was_hr = False
        else:
            prev_was_hr = True
        continue
    else:
        prev_was_hr = False

    # Headings
    heading_match = re.match(r'^(#{1,3})\s+(.*)', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2).strip()
        doc.add_heading(text, level=level)
        continue

    # Bold text as emphasis (not heading)
    if stripped.startswith('**') and stripped.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.strip('*'))
        run.bold = True
        continue

    # Image — skip (text-only version)
    img_match = re.match(r'!\[.*?\]\((.*?)\)', stripped)
    if img_match:
        continue

    # Figure caption
    if stripped.startswith('**Figure') or stripped.startswith('**Table'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('**', ''))
        run.bold = True
        run.font.size = Pt(10)
        continue

    # Regular paragraph
    # Clean markdown formatting
    text = stripped
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold markers
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic markers
    text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code markers

    p = doc.add_paragraph(text)

# Flush remaining
if in_code_block:
    flush_code()
if in_table:
    flush_table()

# ── Save ────────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f"Saved to {DOCX_PATH}")
